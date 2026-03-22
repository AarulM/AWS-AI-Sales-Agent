import io
import os
from typing import Dict, List
import PyPDF2
import docx
import json

class DocumentProcessor:
    """Process various document formats and extract text"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.json', '.csv']
    
    def process_file(self, file_content: bytes, filename: str) -> Dict:
        """
        Process uploaded file and extract text
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Dict with extracted text and metadata
        """
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        try:
            if file_ext == '.pdf':
                return self._process_pdf(file_content, filename)
            elif file_ext == '.docx':
                return self._process_docx(file_content, filename)
            elif file_ext == '.txt':
                return self._process_txt(file_content, filename)
            elif file_ext == '.json':
                return self._process_json(file_content, filename)
            elif file_ext == '.csv':
                return self._process_csv(file_content, filename)
        except Exception as e:
            raise Exception(f"Error processing {filename}: {str(e)}")
    
    def _process_pdf(self, content: bytes, filename: str) -> Dict:
        """Extract text from PDF"""
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_chunks = []
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                text_chunks.append({
                    'page': page_num + 1,
                    'text': text.strip()
                })
        
        return {
            'filename': filename,
            'type': 'pdf',
            'pages': len(pdf_reader.pages),
            'chunks': text_chunks,
            'full_text': '\n\n'.join([c['text'] for c in text_chunks])
        }
    
    def _process_docx(self, content: bytes, filename: str) -> Dict:
        """Extract text from DOCX"""
        doc_file = io.BytesIO(content)
        doc = docx.Document(doc_file)
        
        text_chunks = []
        for para_num, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text_chunks.append({
                    'paragraph': para_num + 1,
                    'text': paragraph.text.strip()
                })
        
        return {
            'filename': filename,
            'type': 'docx',
            'paragraphs': len(doc.paragraphs),
            'chunks': text_chunks,
            'full_text': '\n\n'.join([c['text'] for c in text_chunks])
        }
    
    def _process_txt(self, content: bytes, filename: str) -> Dict:
        """Extract text from TXT"""
        text = content.decode('utf-8')
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        text_chunks = []
        for i, para in enumerate(paragraphs):
            text_chunks.append({
                'paragraph': i + 1,
                'text': para
            })
        
        return {
            'filename': filename,
            'type': 'txt',
            'paragraphs': len(paragraphs),
            'chunks': text_chunks,
            'full_text': text
        }
    
    def _process_json(self, content: bytes, filename: str) -> Dict:
        """Extract text from JSON"""
        data = json.loads(content.decode('utf-8'))
        
        # Convert JSON to readable text
        text = json.dumps(data, indent=2)
        
        return {
            'filename': filename,
            'type': 'json',
            'data': data,
            'chunks': [{'text': text}],
            'full_text': text
        }
    
    def _process_csv(self, content: bytes, filename: str) -> Dict:
        """Extract text from CSV"""
        import csv
        
        text = content.decode('utf-8')
        reader = csv.DictReader(io.StringIO(text))
        
        rows = list(reader)
        text_chunks = []
        
        for i, row in enumerate(rows):
            row_text = ', '.join([f"{k}: {v}" for k, v in row.items()])
            text_chunks.append({
                'row': i + 1,
                'text': row_text
            })
        
        return {
            'filename': filename,
            'type': 'csv',
            'rows': len(rows),
            'chunks': text_chunks,
            'full_text': '\n'.join([c['text'] for c in text_chunks])
        }
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for better context"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > chunk_size * 0.5:  # Only break if we're past halfway
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1
            
            if chunk.strip():
                chunks.append(chunk.strip())
            
            start = end - overlap
        
        return chunks
    
    def auto_categorize(self, text: str) -> str:
        """Auto-detect category based on content"""
        text_lower = text.lower()
        
        # Service keywords
        if any(word in text_lower for word in ['service', 'offer', 'provide', 'specialize', 'expert']):
            return 'services'
        
        # Pricing keywords
        if any(word in text_lower for word in ['price', 'cost', '$', 'rate', 'fee', 'charge']):
            return 'pricing'
        
        # FAQ keywords
        if any(word in text_lower for word in ['question', 'answer', 'faq', 'how to', 'what is', 'why']):
            return 'faqs'
        
        # Contact keywords
        if any(word in text_lower for word in ['contact', 'phone', 'email', 'address', 'location']):
            return 'contact'
        
        # About keywords
        if any(word in text_lower for word in ['about', 'history', 'mission', 'vision', 'founded']):
            return 'company'
        
        return 'general'
    
    def extract_key_points(self, text: str, max_points: int = 5) -> List[str]:
        """Extract key bullet points from text"""
        # Split into sentences
        sentences = [s.strip() for s in text.replace('\n', '. ').split('.') if s.strip()]
        
        # Filter for meaningful sentences (not too short, not too long)
        key_points = []
        for sentence in sentences[:max_points * 2]:
            if 20 < len(sentence) < 200:
                key_points.append(sentence)
                if len(key_points) >= max_points:
                    break
        
        return key_points
